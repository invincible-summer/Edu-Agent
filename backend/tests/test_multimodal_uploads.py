"""上传文件多模态对齐（multimodal_parser v2）契约。

docx/pptx 内嵌图、PDF 稠密页插图/表格、md data-uri 图统一分型描述
（题目转录 / 图述 / 装饰丢弃），产出 [图|...]/[表|...] 标记块，与教材
图表管线同构、由 Structured Chunker V2 识别为 figure/table 块。
"""
from __future__ import annotations

import asyncio
import io
import unittest
from unittest.mock import patch

from app.core import multimodal_parser as mp
from app.core.structured_chunker import chunk_text_v2


def _docx_with_image(image_png: bytes, text: str = "讲义正文第一段。") -> bytes:
    import docx as docx_lib
    doc = docx_lib.Document()
    doc.add_paragraph(text)
    doc.add_picture(io.BytesIO(image_png))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _tiny_png() -> bytes:
    import fitz
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 40, 40))
    pix.clear_with(200)
    data = pix.tobytes("png")
    return data


def _pdf_dense_with_figure() -> bytes:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    # ASCII 文本层（默认字体不渲染中文）：保证该页判定为稠密文本层
    page.insert_text((72, 90), "Chapter 1 mechanics lecture notes, dense text layer. " * 3)
    img = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 200, 150))
    img.clear_with(120)
    page.insert_image(fitz.Rect(72, 200, 380, 400), pixmap=img)
    raw = doc.tobytes()
    doc.close()
    return raw


class TestOfficeEmbeddedImages(unittest.TestCase):
    def test_docx_image_described_as_figure_block(self):
        raw = _docx_with_image(_tiny_png())

        async def fake_desc(data):
            return "图述\n一张受力示意图，包含斜面与物块。"

        with patch.object(mp, "describe_embedded_image", fake_desc):
            result = asyncio.run(mp.extract_text_async("讲义.docx", raw))
        self.assertIn("[图|文档图片 1]", result.text)
        self.assertIn("图述", result.text)
        self.assertTrue(result.used_ocr)
        # chunker 识别为 figure 块
        chunks = chunk_text_v2(result.text, source="讲义.docx", file_id="d1")
        types = {c.metadata["block_types"][0] for c in chunks}
        self.assertIn("figure", types)

    def test_docx_image_transcription_block(self):
        raw = _docx_with_image(_tiny_png())

        async def fake_desc(data):
            return "题目转录\n一个质量为 2kg 的物块静止在水平面上，求支持力大小。"

        with patch.object(mp, "describe_embedded_image", fake_desc):
            result = asyncio.run(mp.extract_text_async("练习.docx", raw))
        self.assertIn("[图|文档图片 1]", result.text)
        self.assertIn("质量为 2kg", result.text)
        chunks = chunk_text_v2(result.text, source="练习.docx", file_id="d2")
        types = {c.metadata["block_types"][0] for c in chunks}
        self.assertIn("figure", types)

    def test_decoration_images_dropped(self):
        raw = _docx_with_image(_tiny_png())

        async def fake_desc(data):
            return "装饰"

        with patch.object(mp, "describe_embedded_image", fake_desc):
            result = asyncio.run(mp.extract_text_async("讲义.docx", raw))
        self.assertNotIn("[图|", result.text)
        self.assertIn("讲义正文第一段", result.text)


class TestPdfDenseFigureHarvest(unittest.TestCase):
    def test_dense_pdf_figure_described_and_merged(self):
        raw = _pdf_dense_with_figure()

        async def fake_desc(data):
            return "图述：斜面模型示意图。"

        with patch.object(mp, "describe_embedded_image", fake_desc):
            result = asyncio.run(mp.extract_text_async("讲义.pdf", raw))
        self.assertIn("[图|插图]", result.text)
        self.assertIn("图述：斜面模型示意图。", result.text)
        self.assertIn("mechanics", result.text)  # 原文本层保留
        chunks = chunk_text_v2(result.text, source="讲义.pdf", file_id="p1")
        types = {c.metadata["block_types"][0] for c in chunks}
        self.assertIn("figure", types)

    def test_harvest_failure_never_blocks_upload(self):
        raw = _pdf_dense_with_figure()

        async def boom(data):
            raise RuntimeError("vision down")

        with patch.object(mp, "describe_embedded_image", boom):
            result = asyncio.run(mp.extract_text_async("讲义.pdf", raw))
        # 图述失败按"丢弃该图"降级：原文本层完整保留，上传绝不失败
        self.assertIn("mechanics", result.text)
        self.assertNotIn("[图|插图]", result.text)


class TestMarkdownDataUri(unittest.TestCase):
    def test_md_datauri_image_described(self):
        import base64
        b64 = base64.b64encode(_tiny_png()).decode()
        md = f"# 笔记\n\n坐标图如下：\n\n![fig](data:image/png;base64,{b64})\n"

        async def fake_desc(data):
            return "图述：二次函数图像。"

        with patch.object(mp, "describe_embedded_image", fake_desc):
            result = asyncio.run(mp.extract_text_async("笔记.md", md.encode()))
        self.assertIn("[图|内嵌图 1]", result.text)
        self.assertIn("二次函数图像", result.text)

    def test_plain_md_untouched(self):
        md = "# 笔记\n\n纯文本笔记，无图片。"
        result = asyncio.run(mp.extract_text_async("笔记.md", md.encode()))
        self.assertEqual(result.text.strip(), md.strip())
        self.assertEqual(result.media_count, 0)


if __name__ == "__main__":
    unittest.main()
